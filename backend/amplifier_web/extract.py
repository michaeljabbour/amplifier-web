"""
File content extraction for PDFs, DOCX, and other document types.

Provides server-side text extraction so documents can be included
in LLM context as text content.
"""

from __future__ import annotations

import base64
import io
import logging
from typing import Any

logger = logging.getLogger(__name__)


async def extract_text(filename: str, content_b64: str) -> dict[str, Any]:
    """
    Extract text from a file based on its type.

    Args:
        filename: Original filename (used to determine type)
        content_b64: Base64-encoded file content

    Returns:
        Dict with either "text" (extracted content) or "error" (failure message)
    """
    if not content_b64:
        return {"error": "No content provided"}

    try:
        file_bytes = base64.b64decode(content_b64)
    except Exception as e:
        return {"error": f"Invalid base64 content: {e}"}

    filename_lower = filename.lower()

    try:
        if filename_lower.endswith(".pdf"):
            return await _extract_pdf(file_bytes, filename)
        elif filename_lower.endswith(".docx"):
            return await _extract_docx(file_bytes, filename)
        elif filename_lower.endswith(".txt") or filename_lower.endswith(".md"):
            return await _extract_text_file(file_bytes, filename)
        else:
            return {"error": f"Unsupported file type: {filename}"}

    except Exception as e:
        logger.exception(f"Failed to extract text from {filename}")
        return {"error": f"Extraction failed: {e}"}


async def _extract_pdf(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {"error": "PDF extraction not available (PyMuPDF not installed)"}

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []

        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")

        doc.close()

        if not text_parts:
            return {
                "text": "[PDF contains no extractable text - may be scanned/image-based]",
                "warning": "No text found in PDF",
            }

        return {
            "text": "\n\n".join(text_parts),
            "pages": len(text_parts),
        }

    except Exception as e:
        return {"error": f"PDF extraction failed: {e}"}


async def _extract_docx(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "DOCX extraction not available (python-docx not installed)"}

    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        if not text_parts:
            return {
                "text": "[Document contains no extractable text]",
                "warning": "No text found in document",
            }

        return {"text": "\n\n".join(text_parts)}

    except Exception as e:
        return {"error": f"DOCX extraction failed: {e}"}


async def _extract_text_file(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Extract text from plain text files."""
    # Try common encodings
    for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
        try:
            text = file_bytes.decode(encoding)
            return {"text": text}
        except (UnicodeDecodeError, LookupError):
            continue

    return {"error": "Could not decode text file - unknown encoding"}


# Supported file types for reference
SUPPORTED_TYPES = {
    ".pdf": "PDF documents",
    ".docx": "Word documents",
    ".txt": "Plain text files",
    ".md": "Markdown files",
}

IMAGE_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
}


def get_file_type(filename: str) -> str | None:
    """Get the MIME type for an image file, or None if not an image."""
    filename_lower = filename.lower()
    for ext, mime_type in IMAGE_TYPES.items():
        if filename_lower.endswith(ext):
            return mime_type
    return None


def is_supported_document(filename: str) -> bool:
    """Check if a filename is a supported document type."""
    filename_lower = filename.lower()
    return any(filename_lower.endswith(ext) for ext in SUPPORTED_TYPES)


def is_image(filename: str) -> bool:
    """Check if a filename is a supported image type."""
    return get_file_type(filename) is not None
